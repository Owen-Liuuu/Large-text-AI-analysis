"""Generate a Word (.docx) evaluation report with comparison charts.

Produces a professional document with:
  - Cover page: review title, run metadata
  - Executive summary: overall verdict, key findings
  - Step 1-2: search & verification results
  - Step 3-4: per-paper data comparison table
      (Student vs LLM-A vs LLM-B, colour-coded)
  - Bar chart: agreement rates per paper
  - Bar chart: agreement rates per field
  - Appendix: all flags
"""
from __future__ import annotations

import io
from datetime import datetime
from pathlib import Path

import structlog

from lit_inspector.pipeline.schemas import PipelineRunResult
from lit_inspector.steps.table_comparison.schemas import TableComparisonResult

logger = structlog.get_logger(__name__)

# Colour constants (RGB tuples)
_GREEN = (198, 239, 206)
_RED = (255, 199, 206)
_YELLOW = (255, 235, 156)
_GREY = (217, 217, 217)
_WHITE = (255, 255, 255)
_DARK = (44, 62, 80)


def generate_docx_report(
    result: PipelineRunResult,
    output_path: Path | str,
) -> Path:
    """Generate a DOCX evaluation report from a pipeline result.

    Args:
        result: The completed pipeline result.
        output_path: Where to save the .docx file.

    Returns:
        The path to the saved report.
    """
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    # -- Global style tweaks --
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    # ================================================================
    # COVER PAGE
    # ================================================================
    _add_cover_page(doc, result)

    # ================================================================
    # 1. EXECUTIVE SUMMARY
    # ================================================================
    doc.add_heading("1. Executive Summary", level=1)

    si = result.student_input
    summary_items = [
        f"Review: {si.review_title}",
        f"Student ID: {si.student_id}",
        f"Run ID: {result.run_id}",
        f"Papers analysed: {len(si.selected_papers)}",
        f"Completed: {result.completed_at or 'N/A'}",
    ]
    for item in summary_items:
        doc.add_paragraph(item, style="List Bullet")

    # Verdict
    if result.report:
        p = doc.add_paragraph()
        p.add_run("Verdict: ").bold = True
        p.add_run(result.report.summary)

    # ================================================================
    # 2. SEARCH VALIDATION (Step 1)
    # ================================================================
    doc.add_heading("2. Search Validation (Step 1)", level=1)

    if result.search_result:
        sr = result.search_result
        tbl = doc.add_table(rows=4, cols=2, style="Light Grid Accent 1")
        _set_cell(tbl, 0, 0, "Search Query")
        _set_cell(tbl, 0, 1, sr.reconstructed_query[:120])
        _set_cell(tbl, 1, 0, "Reported Results")
        _set_cell(tbl, 1, 1, str(sr.reported_count))
        _set_cell(tbl, 2, 0, "Actual Results (PubMed)")
        _set_cell(tbl, 2, 1, str(sr.actual_count))
        _set_cell(tbl, 3, 0, "Reproducible?")
        _set_cell(tbl, 3, 1, "Yes" if sr.is_reproducible else "No")

    if result.papers_in_search:
        doc.add_heading("Papers Found in Search Results", level=2)
        tbl = doc.add_table(
            rows=1 + len(result.papers_in_search), cols=2,
            style="Light Grid Accent 1",
        )
        _set_cell(tbl, 0, 0, "Paper Title", bold=True)
        _set_cell(tbl, 0, 1, "Found?", bold=True)
        for i, (title, found) in enumerate(result.papers_in_search.items(), 1):
            _set_cell(tbl, i, 0, title[:80])
            cell = tbl.cell(i, 1)
            cell.text = "Yes" if found else "No"
            _shade_cell(cell, _GREEN if found else _RED)

    # ================================================================
    # 3. PAPER VERIFICATION (Step 2)
    # ================================================================
    doc.add_heading("3. Paper Verification (Step 2)", level=1)

    if result.verification_results:
        tbl = doc.add_table(
            rows=1 + len(result.verification_results), cols=4,
            style="Light Grid Accent 1",
        )
        for j, h in enumerate(["Paper", "Status", "Confidence", "CrossRef Title"]):
            _set_cell(tbl, 0, j, h, bold=True)

        for i, vr in enumerate(result.verification_results, 1):
            _set_cell(tbl, i, 0, vr.reference.title[:55])
            status_cell = tbl.cell(i, 1)
            status_cell.text = vr.status.value
            colour = {"verified": _GREEN, "not_found": _RED}.get(
                vr.status.value, _YELLOW
            )
            _shade_cell(status_cell, colour)
            _set_cell(tbl, i, 2, f"{vr.confidence:.0%}")
            cr_title = (vr.matched_metadata or {}).get("title", "")
            _set_cell(tbl, i, 3, cr_title[:55])

    # ================================================================
    # 4. DATA EXTRACTION COMPARISON (Steps 3-4) — THE CORE TABLE
    # ================================================================
    doc.add_heading("4. Data Extraction Comparison (Steps 3-4)", level=1)

    doc.add_paragraph(
        "The table below compares the student's extracted data with "
        "AI-extracted values from each LLM, aligned on a canonical field "
        "schema (one row per logical field, after mapping student headers "
        "like 'N', 'EFT/EAT' to canonical names like 'sample_size', "
        "'eat_or_eft_t1dm'). Colours: Green = MATCH, Yellow = PARTIAL / "
        "NEEDS REVIEW, Red = DIFF, Grey = missing on one side."
    )

    # Identify distinct extractor names — prefer the authoritative list from
    # the pipeline run (populated even if some extractors failed); fall back to
    # scanning extracted_tables for backward compatibility.
    extractor_ids: list[str] = [
        e for e in result.extractor_ids if e != "student"
    ] if result.extractor_ids else []
    if not extractor_ids:
        for et in result.extracted_tables:
            if et.extractor_id not in extractor_ids and et.extractor_id != "student":
                extractor_ids.append(et.extractor_id)

    # Build a short display name for each extractor column header.
    # "qwen-plus-extractor" → "Qwen-Plus", "gemini-2.5-flash-extractor" → "Gemini-2.5"
    def _short_name(eid: str) -> str:
        parts = eid.replace("-extractor", "").split("-")
        # Keep up to 2 parts for readability
        return "-".join(p.capitalize() for p in parts[:2])

    if result.comparison_results:
        for cr in result.comparison_results:
            # Find a human-friendly title for the paper
            paper_label = cr.paper_id
            for ref in result.student_input.selected_papers:
                norm_ref = (ref.doi or "").strip().lower()
                norm_cr = cr.paper_id.strip().lower()
                for prefix in ("https://doi.org/", "http://doi.org/", "doi:"):
                    norm_ref = norm_ref.removeprefix(prefix)
                    norm_cr = norm_cr.removeprefix(prefix)
                if norm_ref and norm_ref == norm_cr:
                    # Use first 70 chars of the title
                    paper_label = ref.title[:70]
                    break

            doc.add_heading(
                f"Paper: {paper_label} — Agreement {cr.agreement_rate:.0%}",
                level=2,
            )

            # Columns: Field | Student | <LLM-A> | <LLM-B> ... | Evidence | Status
            col_headers = ["Field", "Student"]
            for eid in extractor_ids:
                col_headers.append(_short_name(eid))
            col_headers.extend(["Evidence", "Status"])
            n_cols = len(col_headers)

            if not cr.field_diffs:
                doc.add_paragraph("(No fields to compare for this paper.)")
                doc.add_paragraph()
                continue

            tbl = doc.add_table(
                rows=1 + len(cr.field_diffs), cols=n_cols,
                style="Light Grid Accent 1",
            )
            for j, h in enumerate(col_headers):
                _set_cell(tbl, 0, j, h, bold=True)

            for i, diff in enumerate(cr.field_diffs, 1):
                # Field column: canonical name on line 1, raw aliases on line 2
                alias_parts = []
                if diff.student_raw_name and diff.student_raw_name != diff.field_name:
                    alias_parts.append(f"student: {diff.student_raw_name}")
                model_raws = [
                    r for r in diff.model_raw_names if r != diff.field_name
                ]
                if model_raws:
                    alias_parts.append("model: " + ", ".join(model_raws[:2]))
                field_label = diff.field_name
                if alias_parts:
                    field_label += "\n(" + "; ".join(alias_parts) + ")"
                _set_cell(tbl, i, 0, field_label)

                # Student value: raw above, normalised below when they differ
                s_raw = _fmt(diff.student_value)
                s_norm = _fmt(diff.student_value_normalized)
                if s_norm and s_norm != s_raw and s_norm != "N/A":
                    _set_cell(tbl, i, 1, f"{s_raw}\n→ {s_norm}")
                else:
                    _set_cell(tbl, i, 1, s_raw)

                # AI values — one column per extractor, in the order the
                # pipeline recorded extractor_ids. We aligned FieldDiff so
                # model_values[j] corresponds to extractor j.
                for j, _eid in enumerate(extractor_ids):
                    if j < len(diff.model_values):
                        ai_raw = _fmt(diff.model_values[j])
                        ai_norm = (
                            _fmt(diff.model_values_normalized[j])
                            if j < len(diff.model_values_normalized)
                            else "N/A"
                        )
                        if ai_norm and ai_norm != ai_raw and ai_norm != "N/A":
                            cell_text = f"{ai_raw}\n→ {ai_norm}"
                        else:
                            cell_text = ai_raw
                    else:
                        cell_text = "—"
                    tbl.cell(i, 2 + j).text = cell_text

                # Evidence column: first non-empty snippet (student first, then model)
                evidence_text = ""
                if diff.student_evidence:
                    evidence_text = f"[student] {diff.student_evidence}"
                else:
                    for snip in diff.model_evidence:
                        if snip:
                            evidence_text = f"[llm] {snip}"
                            break
                if not evidence_text and diff.source_type:
                    evidence_text = f"(source: {diff.source_type})"
                evidence_cell_idx = n_cols - 2
                _set_cell(tbl, i, evidence_cell_idx, _trunc(evidence_text, 120))

                # Status column — render from the rich FieldStatus enum.
                status_cell = tbl.cell(i, n_cols - 1)
                status_val = getattr(diff.status, "value", str(diff.status))
                if status_val == "match":
                    status_cell.text = "✓ MATCH"
                    _shade_cell(status_cell, _GREEN)
                elif status_val == "partial_match":
                    status_cell.text = "~ PARTIAL"
                    _shade_cell(status_cell, _YELLOW)
                elif status_val == "missing_model":
                    status_cell.text = "MODEL N/A"
                    _shade_cell(status_cell, _GREY)
                elif status_val == "missing_student":
                    status_cell.text = "STUDENT N/A"
                    _shade_cell(status_cell, _GREY)
                elif status_val == "not_comparable":
                    status_cell.text = "N/A"
                    _shade_cell(status_cell, _GREY)
                elif status_val == "needs_review":
                    status_cell.text = "? REVIEW"
                    _shade_cell(status_cell, _YELLOW)
                else:  # diff
                    status_cell.text = "✗ DIFF"
                    _shade_cell(status_cell, _RED)

            doc.add_paragraph()  # spacing

    # ================================================================
    # 5. CHARTS
    # ================================================================
    doc.add_heading("5. Visualisation", level=1)

    # 5a. Agreement rate per paper (bar chart)
    if result.comparison_results:
        doc.add_heading("Agreement Rate per Paper", level=2)
        chart_bytes = _chart_agreement_per_paper(result.comparison_results)
        if chart_bytes:
            doc.add_picture(io.BytesIO(chart_bytes), width=Inches(5.8))

    # 5b. Agreement rate per field (averaged across papers)
    if result.comparison_results:
        doc.add_heading("Agreement Rate per Field", level=2)
        chart_bytes = _chart_agreement_per_field(result.comparison_results)
        if chart_bytes:
            doc.add_picture(io.BytesIO(chart_bytes), width=Inches(5.8))

    # ================================================================
    # 6. FLAGS APPENDIX
    # ================================================================
    doc.add_heading("6. All Flags", level=1)

    if result.all_flags:
        tbl = doc.add_table(
            rows=1 + len(result.all_flags), cols=4,
            style="Light Grid Accent 1",
        )
        for j, h in enumerate(["Severity", "Step", "Code", "Message"]):
            _set_cell(tbl, 0, j, h, bold=True)
        for i, flag in enumerate(result.all_flags, 1):
            sev_cell = tbl.cell(i, 0)
            sev_cell.text = flag.severity.value.upper()
            colour = {"error": _RED, "warning": _YELLOW, "info": _GREEN}.get(
                flag.severity.value, _WHITE
            )
            _shade_cell(sev_cell, colour)
            _set_cell(tbl, i, 1, flag.step.value)
            _set_cell(tbl, i, 2, flag.code)
            _set_cell(tbl, i, 3, flag.message[:100])
    else:
        doc.add_paragraph("No flags raised.")

    # ================================================================
    # SAVE
    # ================================================================
    doc.save(str(output_path))
    logger.info("docx_report_saved", path=str(output_path))
    return output_path


# ======================================================================
# Chart helpers (matplotlib)
# ======================================================================


def _chart_agreement_per_paper(
    comparisons: list[TableComparisonResult],
) -> bytes | None:
    """Generate a horizontal bar chart: agreement rate per paper."""
    try:
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt
    except ImportError:
        logger.warning("matplotlib_not_installed")
        return None

    labels = []
    rates = []
    colours = []
    for cr in comparisons:
        # Short label: first author + DOI suffix
        short = cr.paper_id
        if "/" in short:
            short = short.split("/")[-1]
        if len(short) > 25:
            short = short[:22] + "..."
        labels.append(short)
        rates.append(cr.agreement_rate * 100)
        if cr.agreement_rate >= 0.7:
            colours.append("#27ae60")
        elif cr.agreement_rate >= 0.4:
            colours.append("#f39c12")
        else:
            colours.append("#e74c3c")

    fig, ax = plt.subplots(figsize=(8, max(3, len(labels) * 0.5)))
    y_pos = range(len(labels))
    bars = ax.barh(y_pos, rates, color=colours, edgecolor="white", height=0.6)
    ax.set_yticks(y_pos)
    ax.set_yticklabels(labels, fontsize=8)
    ax.set_xlabel("Agreement Rate (%)")
    ax.set_xlim(0, 105)
    ax.set_title("Student vs AI: Agreement Rate per Paper")
    ax.axvline(x=70, color="#27ae60", linestyle="--", alpha=0.5, label="70% threshold")
    ax.legend(fontsize=7)

    # Add percentage labels on bars
    for bar, rate in zip(bars, rates):
        ax.text(
            bar.get_width() + 1, bar.get_y() + bar.get_height() / 2,
            f"{rate:.0f}%", va="center", fontsize=8,
        )

    ax.invert_yaxis()
    plt.tight_layout()

    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150)
    plt.close(fig)
    return buf.getvalue()


def _chart_agreement_per_field(
    comparisons: list[TableComparisonResult],
) -> bytes | None:
    """Generate a vertical bar chart: agreement rate per field (averaged)."""
    try:
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt
    except ImportError:
        return None

    # Aggregate: for each field_name, count match / total across papers
    field_stats: dict[str, list[bool]] = {}
    for cr in comparisons:
        for diff in cr.field_diffs:
            field_stats.setdefault(diff.field_name, []).append(diff.is_consistent)

    if not field_stats:
        return None

    labels = list(field_stats.keys())
    rates = [
        sum(v) / len(v) * 100 if v else 0
        for v in field_stats.values()
    ]
    colours = [
        "#27ae60" if r >= 70 else "#f39c12" if r >= 40 else "#e74c3c"
        for r in rates
    ]

    fig, ax = plt.subplots(figsize=(max(6, len(labels) * 0.6), 4))
    x_pos = range(len(labels))
    bars = ax.bar(x_pos, rates, color=colours, edgecolor="white", width=0.6)
    ax.set_xticks(x_pos)
    ax.set_xticklabels(labels, rotation=45, ha="right", fontsize=8)
    ax.set_ylabel("Agreement Rate (%)")
    ax.set_ylim(0, 110)
    ax.set_title("Student vs AI: Agreement Rate per Field")
    ax.axhline(y=70, color="#27ae60", linestyle="--", alpha=0.5)

    for bar, rate in zip(bars, rates):
        ax.text(
            bar.get_x() + bar.get_width() / 2, bar.get_height() + 2,
            f"{rate:.0f}%", ha="center", fontsize=7,
        )

    plt.tight_layout()
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150)
    plt.close(fig)
    return buf.getvalue()


# ======================================================================
# Document helpers
# ======================================================================


def _add_cover_page(doc, result: PipelineRunResult) -> None:
    """Add a professional cover page."""
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    # Spacer
    for _ in range(4):
        doc.add_paragraph()

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Literature Integrity Detection Report")
    run.font.size = Pt(26)
    run.bold = True

    # Subtitle
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(result.student_input.review_title)
    run.font.size = Pt(14)
    run.italic = True

    doc.add_paragraph()

    # Metadata
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Student: {result.student_input.student_id}\n").font.size = Pt(11)
    meta.add_run(f"Run ID: {result.run_id}\n").font.size = Pt(11)
    meta.add_run(
        f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}\n"
    ).font.size = Pt(11)

    # Extractors used — prefer authoritative list from pipeline run
    eids = list(result.extractor_ids) if result.extractor_ids else []
    if not eids:
        for et in result.extracted_tables:
            if et.extractor_id not in eids:
                eids.append(et.extractor_id)
    if eids:
        meta.add_run(f"LLM Extractors: {', '.join(eids)}\n").font.size = Pt(11)

    doc.add_page_break()


def _set_cell(
    table, row: int, col: int, text: str, bold: bool = False
) -> None:
    """Set text in a table cell."""
    from docx.shared import Pt

    cell = table.cell(row, col)
    cell.text = text
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(8)
            if bold:
                run.bold = True


def _shade_cell(cell, rgb: tuple[int, int, int]) -> None:
    """Apply a background colour to a table cell."""
    from docx.oxml.ns import qn

    shading = cell._element.get_or_add_tcPr()
    shading_elm = shading.find(qn("w:shd"))
    if shading_elm is None:
        from docx.oxml import OxmlElement
        shading_elm = OxmlElement("w:shd")
        shading.append(shading_elm)
    hex_colour = f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"
    shading_elm.set(qn("w:fill"), hex_colour)
    shading_elm.set(qn("w:val"), "clear")


def _fmt(value: object) -> str:
    """Format a value for display in a table cell."""
    if value is None:
        return "N/A"
    s = str(value)
    if len(s) > 35:
        s = s[:32] + "..."
    return s


def _trunc(text: str, limit: int) -> str:
    """Hard-truncate with ellipsis for evidence cells."""
    if not text:
        return ""
    if len(text) <= limit:
        return text
    return text[: max(0, limit - 3)] + "..."
